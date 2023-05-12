import os
import docx
#查找所有的文档
def read_folder(path):
    # 指定文件夹路径
    folder_path = path
    list_docs=[]
    # 获取文件夹中的所有文件
    files = os.listdir(folder_path)

    # 遍历所有文件，输出带有文件后缀的文件名
    for file in files:
        if os.path.isfile(os.path.join(folder_path, file)):
            #print(file)
            list_docs.append(file)
    return list_docs


#查找文档中的所有内容，并每个15个单元格输出一个列表，后续可以增加一个输入参数：i，来定义多少个单元格一分割
def doc2list(docname):
    # 读取docx文件
    #doc = docx.Document('word_20cent\PRC-0009430 变桨柜安装_F.docx')
    doc=docx.Document(docname)
    content_list = []#每20个单元格存入一个list（去重）
    temp_content = []#用来循环每20个单元格的临时list
    listdoc=[]#文档的名称
    for table in doc.tables:
        i=0
        for row in table.rows:

            for cell in row.cells:

                #print(cell.text)
                temp_content.append(cell.text)
                i+=1
                if i == 25 or (table == doc.tables[-1] and row == table.rows[-1] and cell == row.cells[-1]):
                    new_list = [item for item in temp_content if item.strip()]
                    if len(new_list) != 0:
                        content_list.append(new_list)
                    temp_content = []
                    i = 0
        # 输出结果
    for index, content in enumerate(content_list):
        print(f'{docname}'+'的文档中的'+f'第{index + 1}个列表')

        listdoc.append(content)
        #print(content)

    return(listdoc)


def list_submit(folder_path):
    list_all=[]
#函数执行段

#(read_folder("word_20cent"))
    for i in range(0,len(read_folder(folder_path))):
            list_all.append(doc2list(folder_path+'\\'+read_folder(folder_path)[i]))
    print(list_all)
    return(list_all)

##最终执行
folder_path='world_40cent/40%'

list_submit('world_40cent/40%')

